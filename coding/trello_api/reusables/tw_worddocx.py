import arrow
import docx
import datetime  # comes with normal python so no need to install
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# REPORT CREATION METHODS
# Get todays date
def getTodayDate():
    today = arrow.now()
    newtoday = today.strftime("%d/%m/%Y")
    newtoday = datetime.strptime(newtoday, "%d/%m/%Y").date()
    return newtoday


# Convert full datetime format to arrow datetime
def convertFullDatetimetoDatetime(date, date_format):
    if date is not None:
        # print(f"Starting format is: {date}")
        # print(type(date))
        try:
            new_date = arrow.get(date)
            # print(f"ARROW formatted DATETIME: {new_date}")
            # print(type(new_date))
            return new_date
        except Exception:
            print(f"Something went wrong converting {date}!")
    else:
        print(f"None value was supplied. Moving onto next date...")


# Convert full datetime format to string for report
def convertFullDatetimetoString(date, date_format):
    if date is not None:
        # print(f"Starting format is: {date}")
        # print(type(date))
        try:
            new_date = arrow.get(date)
            # print(f"ARROW formatted DATETIME: {new_date}")
            # print(type(new_date))
            new_date = new_date.strftime(date_format)
            # print(f"DATETIME formatted to STRING: {new_date}")
            # print(type(new_date))
            return new_date
        except Exception:
            print(f"Something went wrong converting {date}!")
    else:
        print(f"None value was supplied. Moving onto next date...")


# Is date 1 greater than date 2?
def isDateGreaterThanOtherDate(date1, date2):
    if date1 > date2:
        print(f"{date1} is greater than {date2}")
        return True
    else:
        print(f"{date1} is NOT greater than {date2}")
        return False


# Intersperse method is used to add a "/" in between card members if there is more than one
def intersperse(lst, item):
    result = [item] * (len(lst) * 2 - 1)
    result[0::2] = lst
    return result


# Can change from portrait to landscape and vice versa, just depends what it currently is
def change_orientation(document):
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


# Sets cell border shading and line thickness
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# Applies styling to paragraph/runs (sentences)
def applyStyle(doc, new_style_name, fontname, fontsize, color):
    # Add Header, title, paragraph styles
    obj_styles = doc.styles
    obj_charstyle = obj_styles.add_style(new_style_name, WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(fontsize)
    obj_font.name = fontname
    obj_font.color.rgb = RGBColor.from_string(color)


def generateWordDoc(full_team_cards, rec_comp_cards, full_app_cards, full_in_cards, templateLoc, outputLoc):
    try:
        print("###################Generating Report...###################")
        # Open template
        doc = docx.Document(templateLoc)

        # Add New Style
        applyStyle(doc, 'TestReportTitle', 'Calibri', 24, '00B0F0')
        applyStyle(doc, 'Author', 'Calibri', 14, '00B0F0')
        applyStyle(doc, 'SubHeading', 'Calibri', 16, 'FF1C60')

        print("Creating Page 1...")
        # Add text to the document
        newtoday = getTodayDate()
        newtoday = convertFullDatetimetoString(newtoday, '%d/%m/%Y')
        title = doc.add_paragraph()
        author = doc.add_paragraph()
        toc = doc.add_paragraph()
        titlerun = title.add_run(f'Internal Testing Report - {newtoday}', style='TestReportTitle').bold = True
        authorrun = author.add_run(f'Author: Colin Denny', style='Author').bold = True
        # Align text
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        author.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        contents = doc.add_paragraph()

        # Add a page break
        doc.add_page_break()

        print("Creating Page 2 - Summary In Progress...")
        # Add Summary of IP Projects
        doc.add_heading('Summary of Test Projects In Progress', 1)

        print("Creating Page 2 - Summary In Progress - Table Header...")
        ipTable = doc.add_table(rows=1, cols=5)
        ipTable.style = 'Light List Accent 1'
        ipTable.allow_autofit = False
        hdr_Cells = ipTable.rows[0].cells
        hdr_Cells[0].text = 'Name of System/ Project'
        hdr_Cells[1].text = 'Prep & Execution Status'
        hdr_Cells[2].text = 'Project Completion Status'
        hdr_Cells[3].text = 'Resource Allocation'
        hdr_Cells[4].text = 'Target End Date'
        for cell in hdr_Cells:
            set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                            start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
            hdrparagraph = cell.paragraphs[0]
            hdrrun = hdrparagraph.runs
            hdrfont = hdrrun[0].font
            hdrfont.size = Pt(10)
            hdrfont.bold = False
            hdrfont.color.rgb = RGBColor.from_string('FFFFFF')
        print("Creating Page 2 - Summary In Progress - Table Body...")
        for cn, cl, cm, cs, cd, cc, cla, c in full_team_cards:
            row_Cells = ipTable.add_row().cells
            row_Cells[0].text = cn
            row_Cells[1].text = ""
            row_Cells[2].text = ""
            row_Cells[3].text = intersperse(cm, "/")
            row_Cells[4].text = str(cd)
            if row_Cells[4].text != "N/A":
                row_Cells[4].text = cd.strftime("%d/%m/%y")
            for cell in row_Cells:
                set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
                rowparagraph = cell.paragraphs[0]
                rowrun = rowparagraph.runs
                rowfont = rowrun[0].font
                rowfont.bold = False

        # Add a page break
        doc.add_page_break()

        print("Creating Page 3 - Summary Recently Complete...")
        # Add Summary of RC Projects
        doc.add_heading('Summary of Test Projects Recently Complete', 1)

        print("Creating Page 3 - Summary Recently Complete - Table Header...")
        rcTable = doc.add_table(rows=1, cols=4)
        rcTable.style = 'Light List Accent 1'
        rcTable.allow_autofit = False
        hdr_Cells = rcTable.rows[0].cells
        hdr_Cells[0].text = 'Name of System/ Project'
        hdr_Cells[1].text = 'Testing Status'
        hdr_Cells[2].text = 'Project Status'
        hdr_Cells[3].text = 'Notes'
        for cell in hdr_Cells:
            set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                            start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
            hdrparagraph = cell.paragraphs[0]
            hdrrun = hdrparagraph.runs
            hdrfont = hdrrun[0].font
            hdrfont.size = Pt(10)
            hdrfont.bold = False
            hdrfont.color.rgb = RGBColor.from_string('FFFFFF')
        try:
            print("Creating Page 3 - Summary Recently Complete - Table Body...")
            if len(rec_comp_cards) > 0:
                for cn, cl, cm, cs, cd, cc, cla, c in rec_comp_cards:
                    row_Cells = rcTable.add_row().cells
                    row_Cells[0].text = cn
                    row_Cells[1].text = "Green"
                    row_Cells[2].text = "Green"
                    row_Cells[3].text = ""
                    for cell in row_Cells:
                        set_cell_border(cell,
                                        top={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                        bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                        start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                        insideH={"sz": 8, "val": "single", "color": "#5b9bd5",
                                                 "space": "0"})
                        rowparagraph = cell.paragraphs[0]
                        rowrun = rowparagraph.runs
                        rowfont = rowrun[0].font
                        rowfont.bold = False
            else:
                print("There are no 'Recently Complete' cards so will add empty table with N/A's.")
                row_Cells = rcTable.add_row().cells
                row_Cells[0].text = "N/A"
                row_Cells[1].text = "N/A"
                row_Cells[2].text = "N/A"
                row_Cells[3].text = "N/A"
                for cell in row_Cells:
                    set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                    bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                    start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                    insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
                    rowparagraph = cell.paragraphs[0]
                    rowrun = rowparagraph.runs
                    rowfont = rowrun[0].font
                    rowfont.bold = False
        except Exception:
            print("Something went wrong trying to add recently complete table")

        # Add a page break
        doc.add_page_break()

        print("Creating Page 4 - Project Commentary...")
        # Add Summary of Commentary
        doc.add_heading('Project Commentary', 1)
        doc.add_heading("Projects In Progress", level=2)

        ordered = "5"
        unordered = "1"
        for ftcard in full_team_cards:
            doc.add_heading(ftcard[0], 3)
            listTodayComment = ftcard[7]
            pip = doc.add_paragraph(listTodayComment, style='List Bullet 2')

        blankpara = doc.add_paragraph()

        doc.add_heading("Projects Recently Complete", level=2)
        if len(rec_comp_cards) > 0:
            for reccard in rec_comp_cards:
                doc.add_heading(reccard[0], 3)
                listTodayComment = reccard[7]
                pip = doc.add_paragraph(listTodayComment, style='List Bullet 2')
        else:
            doc.add_paragraph("There are no recently complete projects.", style='List Bullet 2')

        # Add a page break
        doc.add_page_break()

        print("Creating Page 5 - Future Approved/Imminent...")
        # Add Summary of Commentary
        doc.add_heading('Future Testing Activities', 1)
        doc.add_heading('Approved / Imminent', 2)

        appTable = doc.add_table(rows=1, cols=4)
        appTable.style = 'Light List Accent 1'
        appTable.allow_autofit = False
        hdr_Cells = appTable.rows[0].cells
        hdr_Cells[0].text = 'Name of System/ Project'
        hdr_Cells[1].text = 'Delivery into SIT'
        hdr_Cells[2].text = 'Exit SIT'
        hdr_Cells[3].text = 'Comments'
        for cell in hdr_Cells:
            set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                            start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
            hdrparagraph = cell.paragraphs[0]
            hdrrun = hdrparagraph.runs
            hdrfont = hdrrun[0].font
            hdrfont.size = Pt(10)
            hdrfont.bold = False
            hdrfont.color.rgb = RGBColor.from_string('FFFFFF')
        for cn, cl, cm, cs, cd, cc, cla, c in full_app_cards:
            row_Cells = appTable.add_row().cells
            row_Cells[0].text = cn
            row_Cells[1].text = "TBC"
            row_Cells[2].text = "TBC"
            row_Cells[3].text = cl
            if cs != "N/A":
                row_Cells[1].text = convertFullDatetimetoString(cs, "%d/%m/%y")
            if cd != "N/A":
                row_Cells[2].text = convertFullDatetimetoString(cd, "%d/%m/%y")
            for cell in row_Cells:
                set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
                rowparagraph = cell.paragraphs[0]
                rowrun = rowparagraph.runs
                rowfont = rowrun[0].font
                rowfont.bold = False
        # Add a page break
        doc.add_page_break()

        print("Creating Page 6 - Future Incoming...")
        doc.add_heading('Incoming / Probable / Potential', 2)

        inTable = doc.add_table(rows=1, cols=4)
        inTable.style = 'Light List Accent 1'
        inTable.allow_autofit = False
        hdr_Cells = inTable.rows[0].cells
        hdr_Cells[0].text = 'Name of System/ Project'
        hdr_Cells[1].text = 'Delivery into SIT'
        hdr_Cells[2].text = 'Exit SIT'
        hdr_Cells[3].text = 'Comments'
        for cell in hdr_Cells:
            set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                            start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                            insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
            hdrparagraph = cell.paragraphs[0]
            hdrrun = hdrparagraph.runs
            hdrfont = hdrrun[0].font
            hdrfont.size = Pt(10)
            hdrfont.bold = False
            hdrfont.color.rgb = RGBColor.from_string('FFFFFF')
        for cn, cl, cm, cs, cd, cc, cla, c in full_in_cards:
            row_Cells = inTable.add_row().cells
            row_Cells[0].text = cn
            row_Cells[1].text = "TBC"
            row_Cells[2].text = "TBC"
            row_Cells[3].text = cl
            if cs != "N/A":
                row_Cells[1].text = convertFullDatetimetoString(cs, "%d/%m/%y")
            if cd != "N/A":
                row_Cells[2].text = convertFullDatetimetoString(cd, "%d/%m/%y")
            for cell in row_Cells:
                set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                bottom={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"},
                                start={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                end={"sz": 8, "val": "single", "color": "#000000", "space": "0"},
                                insideH={"sz": 8, "val": "single", "color": "#5b9bd5", "space": "0"})
                rowparagraph = cell.paragraphs[0]
                rowrun = rowparagraph.runs
                rowfont = rowrun[0].font
                rowfont.bold = False
        # Add a page break
        doc.add_page_break()

        print("Creating Page 7 - Other News...")
        # Add Summary of Commentary
        doc.add_heading('Other News', 1)
        doc.add_heading('Approaching Holidays', 2)
        ah = doc.add_paragraph('Holiday Info goes here', style='List Bullet 2')
        doc.add_heading('Other Leave / Out of Office', 2)
        ol = doc.add_paragraph('All – Working from home in line with guidance.', style='List Bullet 2')
        doc.add_heading('Annual Probationary Reviews', 2)
        apr = doc.add_paragraph('N/A', style='List Bullet 2')

        # Save new file
        doc.save(outputLoc + "Testing Weekly Report " + newtoday.replace("/", "") + ".docx")
        print("Your report has been created")
    except Exception:
        print("WOAH something went wrong. Check the latest print statement to find the last action before failure!!!!!")
